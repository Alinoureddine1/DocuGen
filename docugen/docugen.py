# Copyright (C) 2024 DocuGen
# This file is part of DocuGen.
#
# DocuGen is free software: you can redistribute it and/or modify
# it under the terms of the GNU General Public License as published by
# the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU General Public License for more details.
#
# You should have received a copy of the GNU General Public License
# along with this program.  If not, see <https://www.gnu.org/licenses/>.

import os
import re
import random
import argparse
from typing import List, Optional
from tqdm import tqdm
import logging
import wikipedia
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
import names
from .text_processor import paraphrase

# Try to import reportlab, but don't fail if it's not available
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("reportlab not installed. PDF generation will be unavailable.")

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def clean_title(title: str, max_length: int = 50) -> str:
    """Clean and format the title."""
    # Remove content within parentheses, brackets, and curly braces
    cleaned = re.sub(r'\([^)]*\)|\[[^]]*\]|\{[^}]*\}', '', title)
    # Remove special characters including backticks and equals signs
    cleaned = re.sub(r'[(){}\[\]`=]', '', cleaned)
    # Remove external links section and references
    cleaned = re.sub(r'== external links ==.*', '', cleaned, flags=re.IGNORECASE | re.DOTALL)
    cleaned = re.sub(r'== references ==.*', '', cleaned, flags=re.IGNORECASE | re.DOTALL)
    # Replace multiple spaces with a single space
    cleaned = re.sub(r'\s+', ' ', cleaned)
    # Remove any non-alphanumeric characters from the start and end
    cleaned = cleaned.strip()
    # Capitalize the first letter of each word
    cleaned = ' '.join(word.capitalize() for word in cleaned.split())
    # Limit to specified max_length and strip again to remove any trailing space
    cleaned = cleaned[:max_length].strip()
    # If the title is now empty, use a default title
    if not cleaned:
        cleaned = "Untitled Document"
    return cleaned

def get_wikipedia_page(title: Optional[str] = None) -> wikipedia.WikipediaPage:
    """Retrieve a Wikipedia page."""
    try:
        if not title:
            title = wikipedia.random(1)
        page = wikipedia.page(title, auto_suggest=False)
        logging.info(f"Retrieved page: {page.title} (URL: {page.url})")
        return page
    except wikipedia.DisambiguationError as e:
        title = random.choice(e.options)
        page = wikipedia.page(title, auto_suggest=False)
        logging.info(f"Disambiguation: using '{page.title}' (URL: {page.url})")
        return page
    except Exception as e:
        logging.error(f"Error retrieving Wikipedia page: {e}")
        raise

def chunk_content(content: str, chunk_size: int) -> List[str]:
    """Split content into chunks of specified sentence count."""
    sentences = sent_tokenize(content)
    chunks = []
    current_chunk = []
    
    for sentence in sentences:
        current_chunk.append(sentence)
        if len(current_chunk) >= chunk_size:
            chunks.append(" ".join(current_chunk))
            current_chunk = []
    
    if current_chunk:
        chunks.append(" ".join(current_chunk))
    
    return chunks

def create_docx(text: str, title: str, folder: str, class_name: str, doc_no: int, file_title: str) -> None:
    """Create a Word document."""
    file_path = f"{folder}/{file_title}.docx"
    
    try:
        document = Document()
        
        # Name and Class
        for content in [names.get_full_name(), class_name]:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(content)
            run.font.size = Pt(12)
            run.font.name = "Times New Roman"
        
        # Heading
        paragraph = document.add_paragraph()
        run = paragraph.add_run(title)
        run.font.size = Pt(16)
        run.font.name = "Times New Roman"
        run.bold = True
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Content
        document.add_paragraph(text)

        document.save(file_path)
        logging.info(f"Created document: {file_path}")
    except Exception as e:
        logging.error(f"Failed to create DOCX {file_path}: {e}")
        raise

def create_pdf(text: str, title: str, folder: str, class_name: str, doc_no: int, file_title: str) -> None:
    """Create a PDF document."""
    file_path = f"{folder}/{file_title}.pdf"
    
    doc = SimpleDocTemplate(file_path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    try:
        # Name and Class
        story.append(Paragraph(names.get_full_name(), styles['Normal']))
        story.append(Spacer(1, 12))
        story.append(Paragraph(class_name, styles['Normal']))
        story.append(Spacer(1, 12))

        # Heading
        title_style = ParagraphStyle(name='Title', parent=styles['Heading1'], alignment=TA_CENTER)
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 12))

        # Content
        story.append(Paragraph(text, styles['BodyText']))

        doc.build(story)
        logging.info(f"Created document: {file_path}")
    except Exception as e:
        logging.error(f"Failed to create PDF {file_path}: {e}")
        raise

def generate_documents(args: argparse.Namespace) -> None:
    """Generate documents based on the provided arguments."""
    amt = args.number
    wiki_title = args.title
    num_sentences = args.sentences
    class_names = args.class_names
    output_format = args.format
    generated = 0
    
    if output_format == "pdf" and not PDF_AVAILABLE:
        logging.error("PDF generation is not available. Please install reportlab or use docx format.")
        return

    # Create a single output folder
    folder = "output"
    os.makedirs(folder, exist_ok=True)
    logging.info(f"Created directory: {folder}")

    with tqdm(total=amt, desc="Generating documents", unit="doc") as pbar:
        while generated < amt:
            page = get_wikipedia_page(wiki_title)
            content = page.content.split("== See also ==")[0].strip()
            chunks = chunk_content(content, num_sentences)
            
            clean_page_title = clean_title(page.title, max_length=30)
            
            for chunk in chunks:
                if generated >= amt:
                    break
                paraphrased_chunk = paraphrase(chunk)
                if len(paraphrased_chunk.split()) < 10:  # Skip if content is too short
                    continue
                doc_title = clean_title(" ".join(word_tokenize(paraphrased_chunk)[:10]), max_length=50)
                
                # Create a short title for the filename
                short_title = clean_title(" ".join(word_tokenize(paraphrased_chunk)[:5]), max_length=20)
                short_title = re.sub(r'\s+', '_', short_title)  # Replace spaces with underscores
                
                # Select a random class name
                class_name = random.choice(class_names)
                
                file_title = f"{generated+1:03d}_{class_name}_{short_title}"
                
                try:
                    if output_format == "docx":
                        create_docx(paraphrased_chunk, doc_title, folder, class_name, doc_no=generated + 1, file_title=file_title)
                    elif PDF_AVAILABLE:
                        create_pdf(paraphrased_chunk, doc_title, folder, class_name, doc_no=generated + 1, file_title=file_title)
                    generated += 1
                    pbar.update(1)
                except Exception as e:
                    logging.warning(f"Failed to create document: {e}")
                    continue

            # If we've exhausted all chunks and still need more documents, get a new page
            if generated < amt:
                wiki_title = None  # This will cause get_wikipedia_page to fetch a random page

    logging.info(f"Successfully generated {generated} documents")

def main() -> None:
    """Main function to parse arguments and initiate document generation."""
    parser = argparse.ArgumentParser(description="Generate study documents from Wikipedia articles")
    parser.add_argument("-t", "--title", help="The Wikipedia page title to use for generating documents", type=str)
    parser.add_argument("-n", "--number", help="The number of documents to generate", default=10, type=int)
    parser.add_argument("-c", "--class_names", help="The class names to use for generating documents", 
                        nargs="+", type=str, default=["CS 101", "CS 202", "CS 303", "CS 404", "CS 505"])
    parser.add_argument("-s", "--sentences", help="The number of sentences to use for each chunk", default=25, type=int)
    parser.add_argument("-f", "--format", help="Output format (docx or pdf)", choices=["docx", "pdf"], default="docx")
    
    args = parser.parse_args()
    
    if os.path.exists("output"):
        import shutil
        shutil.rmtree("output")
        logging.info("Previous output deleted")
    os.makedirs("output", exist_ok=True)
    logging.info("Created output directory")
    
    generate_documents(args)

if __name__ == "__main__":
    main()